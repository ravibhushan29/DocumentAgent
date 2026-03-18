"""PDF (PyMuPDF + OCR fallback) and DOCX parser (Section 7)."""

from __future__ import annotations

import io
from dataclasses import dataclass
from typing import Any

from app.core.exceptions import ParseError

# Minimum chars per page before OCR fallback
MIN_TEXT_CHARS_PER_PAGE = 50


@dataclass
class PageContent:
    """Content for one page (or logical page for DOCX)."""

    page_number: int
    text: str
    metadata: dict[str, Any] | None = None


def parse_pdf(bytes_content: bytes) -> list[PageContent]:
    """Extract text from PDF using PyMuPDF. OCR fallback if page has < 50 chars."""
    try:
        import fitz  # PyMuPDF
    except ImportError as e:
        raise ParseError("pdf", f"PyMuPDF not installed: {e}") from e

    try:
        doc = fitz.open(stream=bytes_content, filetype="pdf")
    except Exception as e:
        raise ParseError("pdf", str(e)) from e

    pages: list[PageContent] = []
    for i in range(len(doc)):
        page = doc.load_page(i)
        text = page.get_text()
        if len(text.strip()) < MIN_TEXT_CHARS_PER_PAGE:
            try:
                import pytesseract
                from PIL import Image
                pix = page.get_pixmap(dpi=150)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                text = pytesseract.image_to_string(img) or text
            except Exception:
                pass  # Keep original text if OCR fails
        pages.append(
            PageContent(
                page_number=i + 1,
                text=text.strip(),
                metadata={"char_count": len(text)},
            )
        )
    doc.close()
    return pages


def parse_docx(bytes_content: bytes) -> list[PageContent]:
    """Extract text from DOCX (paragraphs, headings, tables). Single logical page."""
    try:
        from docx import Document as DocxDocument
    except ImportError as e:
        raise ParseError("docx", f"python-docx not installed: {e}") from e

    try:
        doc = DocxDocument(io.BytesIO(bytes_content))
    except Exception as e:
        raise ParseError("docx", str(e)) from e

    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    text = "\n\n".join(parts)
    return [
        PageContent(
            page_number=1,
            text=text,
            metadata={"paragraphs": len(doc.paragraphs), "tables": len(doc.tables)},
        )
    ]


def parse_file(bytes_content: bytes, file_type: str) -> list[PageContent]:
    """Dispatch by file_type: pdf or docx."""
    if file_type == "pdf":
        return parse_pdf(bytes_content)
    if file_type == "docx":
        return parse_docx(bytes_content)
    raise ParseError(file_type, f"Unsupported type: {file_type}")
